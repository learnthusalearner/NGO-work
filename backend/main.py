from __future__ import annotations

import re
import tempfile
from uuid import uuid4
from contextlib import asynccontextmanager
from datetime import date
from pathlib import Path
from typing import Literal

import sqlite3
from docx import Document
from fastapi import FastAPI, HTTPException, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, field_validator
from starlette.background import BackgroundTask

from database import get_connection, initialize_database, row_to_dict

Role = Literal[
    "Content Writer",
    "Social Media Manager",
    "Field Volunteer",
    "Fundraising Coordinator",
    "Tech Intern",
]
Duration = Literal["1 Month", "2 Months", "3 Months", "6 Months"]
VolunteerStatus = Literal["Active", "Completed"]

ALLOWED_ROLES = set(Role.__args__)
ALLOWED_DURATIONS = set(Duration.__args__)
ALLOWED_STATUSES = set(VolunteerStatus.__args__)
EMAIL_PATTERN = re.compile(r"^[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}$", re.IGNORECASE)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "certificate_template.docx"
DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class VolunteerCreate(BaseModel):
    name: str = Field(..., min_length=2, max_length=120)
    email: str = Field(..., max_length=254)
    role: Role
    duration: Duration

    @field_validator("name")
    @classmethod
    def normalize_name(cls, value: str) -> str:
        normalized = " ".join(value.strip().split())
        if len(normalized) < 2:
            raise ValueError("Name must contain at least two characters.")
        return normalized

    @field_validator("email")
    @classmethod
    def normalize_email(cls, value: str) -> str:
        normalized = value.strip().lower()
        if not EMAIL_PATTERN.match(normalized):
            raise ValueError("Email address is not valid.")
        return normalized


class VolunteerRead(BaseModel):
    id: int
    name: str
    email: str
    role: Role
    duration: Duration
    join_date: str
    status: VolunteerStatus


class StatusUpdate(BaseModel):
    status: VolunteerStatus


@asynccontextmanager
async def lifespan(_: FastAPI):
    initialize_database()
    yield


app = FastAPI(
    title="NayePankh Volunteer Management API",
    description="Volunteer onboarding, roster management, and certificate automation.",
    version="1.0.0",
    lifespan=lifespan,
)

import os

ALLOWED_ORIGINS_ENV = os.environ.get("ALLOWED_ORIGINS")
if ALLOWED_ORIGINS_ENV:
    origins = [o.strip() for o in ALLOWED_ORIGINS_ENV.split(",") if o.strip()]
else:
    origins = [
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:5174",
        "http://127.0.0.1:5174",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
def health_check() -> dict[str, str]:
    return {"status": "ok"}


@app.post(
    "/api/volunteers",
    response_model=VolunteerRead,
    status_code=status.HTTP_201_CREATED,
)
def create_volunteer(payload: VolunteerCreate) -> dict:
    join_date = date.today().isoformat()

    try:
        with get_connection() as connection:
            cursor = connection.execute(
                """
                INSERT INTO volunteers (name, email, role, duration, join_date)
                VALUES (?, ?, ?, ?, ?)
                """,
                (
                    payload.name,
                    payload.email,
                    payload.role,
                    payload.duration,
                    join_date,
                ),
            )
            connection.commit()
            row = connection.execute(
                "SELECT * FROM volunteers WHERE id = ?", (cursor.lastrowid,)
            ).fetchone()
    except sqlite3.IntegrityError as exc:
        if "UNIQUE" in str(exc).upper():
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="A volunteer with this email already exists.",
            ) from exc
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Unable to create volunteer.",
        ) from exc
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Database error while creating volunteer.",
        ) from exc

    created = row_to_dict(row)
    if created is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Volunteer was created but could not be read back.",
        )
    return created


@app.get("/api/volunteers", response_model=list[VolunteerRead])
def list_volunteers(
    search: str | None = Query(default=None, max_length=120),
    role: Role | None = Query(default=None),
) -> list[dict]:
    query = "SELECT * FROM volunteers WHERE 1 = 1"
    params: list[str] = []

    if search and search.strip():
        query += " AND name LIKE ? COLLATE NOCASE"
        params.append(f"%{search.strip()}%")

    if role:
        query += " AND role = ?"
        params.append(role)

    query += " ORDER BY join_date DESC, id DESC"

    try:
        with get_connection() as connection:
            rows = connection.execute(query, params).fetchall()
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Database error while fetching volunteers.",
        ) from exc

    return [dict(row) for row in rows]


@app.patch("/api/volunteers/{volunteer_id}/status", response_model=VolunteerRead)
def update_volunteer_status(
    volunteer_id: int,
    payload: StatusUpdate,
) -> dict:
    try:
        with get_connection() as connection:
            existing = connection.execute(
                "SELECT id FROM volunteers WHERE id = ?", (volunteer_id,)
            ).fetchone()
            if existing is None:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Volunteer not found.",
                )

            connection.execute(
                "UPDATE volunteers SET status = ? WHERE id = ?",
                (payload.status, volunteer_id),
            )
            connection.commit()
            row = connection.execute(
                "SELECT * FROM volunteers WHERE id = ?", (volunteer_id,)
            ).fetchone()
    except HTTPException:
        raise
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Database error while updating volunteer status.",
        ) from exc

    updated = row_to_dict(row)
    if updated is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Volunteer not found.",
        )
    return updated


@app.post("/api/volunteers/{volunteer_id}/generate-certificate")
def generate_certificate(volunteer_id: int) -> FileResponse:
    volunteer = get_volunteer_by_id(volunteer_id)
    template_path = ensure_certificate_template()

    try:
        document = Document(template_path)
        replacements = {
            "{{NAME}}": volunteer["name"],
            "{{ROLE}}": volunteer["role"],
            "{{DURATION}}": volunteer["duration"],
        }
        replace_placeholders(document, replacements)

        output_dir = get_writable_certificate_dir()
        file_stem = safe_filename(volunteer["name"]) or f"volunteer_{volunteer_id}"
        output_path = output_dir / f"{file_stem}_{uuid4().hex}_Certificate.docx"
        document.save(output_path)
    except HTTPException:
        raise
    except (OSError, PermissionError) as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Certificate could not be written to disk.",
        ) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Certificate generation failed.",
        ) from exc

    raw_download_name = f"{file_stem}_Certificate.docx"
    download_name = re.sub(r"\.docx?$", ".docs", raw_download_name, flags=re.IGNORECASE)
    return FileResponse(
        path=output_path,
        media_type=DOCX_MEDIA_TYPE,
        filename=download_name,
        background=BackgroundTask(lambda: output_path.unlink(missing_ok=True)),
    )


def get_volunteer_by_id(volunteer_id: int) -> dict:
    try:
        with get_connection() as connection:
            row = connection.execute(
                "SELECT * FROM volunteers WHERE id = ?", (volunteer_id,)
            ).fetchone()
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Database error while reading volunteer.",
        ) from exc

    volunteer = row_to_dict(row)
    if volunteer is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Volunteer not found.",
        )
    return volunteer


def ensure_certificate_template() -> Path:
    try:
        TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
        if TEMPLATE_PATH.exists():
            return TEMPLATE_PATH

        document = Document()
        document.add_heading("Dummy Volunteer Certificate", level=0)
        document.add_paragraph("This is a dummy certificate file.")
        document.add_paragraph("Volunteer Name: {{NAME}}")
        document.add_paragraph("Volunteer Role: {{ROLE}}")
        document.add_paragraph("Volunteer Duration: {{DURATION}}")
        document.save(TEMPLATE_PATH)
        return TEMPLATE_PATH
    except (OSError, PermissionError) as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Certificate template could not be initialized.",
        ) from exc


def replace_placeholders(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in iter_document_paragraphs(document):
        replace_placeholders_in_paragraph(paragraph, replacements)


def iter_document_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_placeholders_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    original_text = "".join(run.text for run in paragraph.runs)
    if not original_text:
        return

    updated_text = original_text
    for placeholder, value in replacements.items():
        updated_text = updated_text.replace(placeholder, value)

    if updated_text == original_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = updated_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated_text)


def get_writable_certificate_dir() -> Path:
    candidates = [
        Path(tempfile.gettempdir()) / "nayepankh_certificates",
        BASE_DIR / "generated_certificates",
    ]

    for candidate in candidates:
        try:
            candidate.mkdir(parents=True, exist_ok=True)
            probe_path = candidate / ".write_probe"
            probe_path.write_text("ok", encoding="utf-8")
            probe_path.unlink(missing_ok=True)
            return candidate
        except OSError:
            continue

    raise HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail="No writable directory is available for certificate generation.",
    )


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return cleaned.strip("._-")
